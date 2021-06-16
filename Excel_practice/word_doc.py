import docx
d=docx.Document('Test.docx')#document object
print(d)
p=d.paragraphs##paragraph object
print(p)
l=[]
for i in p:
    l.append(i.text)
print(''.join(l))

##new doc
d1=docx.Document()
d1.add_paragraph('Hello')
d1.add_paragraph('This is next paragraph')
d1.save('DemoDocCreation.docx')
p1=d1.paragraphs[0]
p1.add_run('This is a new run')
d1.save('DemoDocCreation.docx')

